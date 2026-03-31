# -*- coding: utf-8 -*-
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FIG = r"D:\AI-agent\openclaw-apri\projects\2025-新能源华南双碳报告\academic\figures"

def sf(r, sz=11, bd=False, cl=None, fn='微软雅黑'):
    r.font.name = fn
    r.font.size = Pt(sz)
    r.font.bold = bd
    if cl: r.font.color.rgb = RGBColor(*cl)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)

def ap(doc, txt, sz=11, bd=False, al=WD_ALIGN_PARAGRAPH.JUSTIFY, ind=False, fn='微软雅黑', cl=None):
    p = doc.add_paragraph()
    p.alignment = al
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if ind: p.paragraph_format.first_line_indent = Pt(22)
    r = p.add_run(txt)
    sf(r, sz, bd, cl, fn)
    return p

def ah(doc, txt, lv=1):
    h = doc.add_heading('', level=lv)
    h.clear()
    r = h.add_run(txt)
    colors = {1:(0,51,102), 2:(31,78,121), 3:(68,68,68)}
    sizes = {1:14, 2:12, 3:11}
    fonts = {1:'黑体', 2:'黑体', 3:'微软雅黑'}
    sf(r, sizes[lv], True, colors[lv], fonts[lv])
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)

def afig(doc, fn, cap, w=5.8):
    fp = FIG + '\\' + fn
    if os.path.exists(fp):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(fp, width=Inches(w))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    r = c.add_run(cap)
    sf(r, 10, False, (100,100,100))

def atab(doc, hdrs, data, cap):
    cp = doc.add_paragraph()
    r = cp.add_run(cap)
    sf(r, 10, False, (100,100,100))
    t = doc.add_table(rows=len(data)+1, cols=len(hdrs))
    t.style = 'Table Grid'
    for j, h in enumerate(hdrs):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].clear()
        rr = cell.paragraphs[0].add_run(h)
        sf(rr, 10, True, (255,255,255), '黑体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E79')
        tcPr.append(shd)
    for i, row in enumerate(data):
        for j, v in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.paragraphs[0].clear()
            rr = cell.paragraphs[0].add_run(str(v))
            sf(rr, 10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6F1')
                tcPr.append(shd)
    doc.add_paragraph()
